import os.path
import pickle
import re
import time
from itertools import chain
from subprocess import run

import docx
import xlsxwriter
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def makeexcel(dir,conten):
    if len(conten) > 0 and len(conten[0]) == 15:
        file = dir+"\\案件客户清单.xlsx"
    else:
        file = dir+"\\法务客户清单.xlsx"
    excel = xlsxwriter.Workbook(file)
    sheet = excel.add_worksheet("sheet1")
    r = 0
    for row in conten:
        row = (*row[9:14],*row[6:8],row[2],re.sub("[0-1\n]*$","",row[-1]))
        c = 0
        for value in row:
            if type(value) is bytes:
                print(pickle.loads(value))
                value = "、".join(i[0] for i in pickle.loads(value))
            sheet.write(r,c, value)
            c+=1
        r+=1
    excel.close()
    os.startfile(file)


def output_paper():
    while True:
        if os.path.exists("data_out.bt"):
            with open("data_out.bt","rb") as t:
                data = pickle.loads(t.read())
            os.remove("data_out.bt")
            break
        else:
            time.sleep(1)
    dir = data[0]
    conten = data[1]
    what = data[2]

    if what == "客户清单":
        makeexcel(dir,conten)
        return
    
    docs = docx.Document()
    p = docs.add_paragraph()
    p.alignment = WD_TABLE_ALIGNMENT.CENTER

    format = p.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(0)

    arun = p.add_run(what)
    arun.font.size = Pt(18)
    arun.font.name = u'黑体'
    arun.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # docs.add_paragraph()

    if re.search("(申请书)|(起诉状)",what) is not None:
        if "申请书" in what:
            k_v = {9: "申请人", 10: "被申请人", 11: "第三人"}
            qingqiu = "仲裁请求"
        else:
            k_v = {9:"原告",10:"被告",11:"第三人"}
            qingqiu = "诉讼请求"
        for i in k_v.keys():
            for j in conten[i]:
                if j[0] == "":
                    continue
                p = docs.add_paragraph()
                p.add_run(k_v[i]+"：").bold = True
                p.add_run(j[0]+"，"+j[1].split("\n")[0])
                for k in j[1].split("\n")[1:]:
                    docs.add_paragraph().add_run(k)

        docs.add_paragraph()
        p = docs.add_paragraph()
        p.add_run("案由：").bold = True
        p.add_run(conten[12][0][0])
        docs.add_paragraph()

        docs.add_paragraph().add_run(qingqiu+"：").bold = True
        n = 0
        for i in conten[3].split("\n"):
            n = n + 1
            docs.add_paragraph(str(n)+". "+i)

        docs.add_paragraph().add_run("事实理由：").bold = True
        if "申请人" in k_v[9]:
            txt = "劳动争议调解仲裁法"
            yuan = ("委","委","裁")
            juzhuangren = "申请人"
        else:
            txt = "民事诉讼法"
            yuan = ("院","院","判")
            juzhuangren = "具状人"
        for i in [*(conten[4].split("\n")),"现为维护自身的权利，{}根据《中华人民共和国{}》及相关规定，向贵{}提起诉讼，请贵{}{}如所请。".format(k_v[9],txt,*yuan),"此致",conten[13][0][0]]:
            docs.add_paragraph().add_run(i)

        p = docs.add_paragraph()
        p.add_run(juzhuangren+"：")
        p.paragraph_format.left_indent = Pt(252)
        p = docs.add_paragraph()
        p.add_run("年    月    日")
        p.alignment = WD_TABLE_ALIGNMENT.RIGHT

        docs.add_paragraph()

        docs.add_paragraph().add_run("附：1. 起诉状副本   份")
        for i in ["2. 证据：","1){}身份证复印件/营业执照复印件、法定代表人身份证明、法定代表人身份证复印件；".format(k_v[9]),"2){}身份证复印件/身份信息/营业执照复印件/企业信用信息打印件；".format(k_v[10])]:
            p = docs.add_paragraph()
            p.add_run(i)
            p.paragraph_format.left_indent = Pt(28)

        ps = docs.paragraphs
        for i in ps[1:]:
            format = i.paragraph_format
            if i not in [ps[-8],ps[-2:]]:
                format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    elif what == "工作日志":
        for i in conten:
            docs.add_paragraph().add_run(i[1].replace("-",".") +"-"+i[-2].replace("-","."))
            docs.add_paragraph().add_run(i[3])
            if i[4] != "":
                docs.add_paragraph().add_run(i[4])
            docs.add_paragraph()

        for i in docs.paragraphs[1:]:
            format = i.paragraph_format
            format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    else:
        other = []
        client = ["","",""]
        k_v = [9, 10,11]
        for i in k_v:
            for j in conten[i]:
                if "委托人" in j:
                    client = j
                else:
                    other.append(j)
        p = docs.add_paragraph()
        p.add_run("答辩人：").bold = True
        p.add_run(client[0] + "，" + client[1].split("\n")[0])
        for i in client[1].split("\n")[1:]:
            docs.add_paragraph().add_run(i)

        docs.add_paragraph()
        docs.add_paragraph().add_run("答辩人与"+"、".join([i[0] for i in other])+conten[12][0][0]+"一案，针对原告/上诉人/申请人的起诉/上诉/仲裁理由答辩如下：")

        docs.add_paragraph().add_run("原告的诉讼请求缺乏事实与法律依据，请求法院驳回原告诉讼请求。理由如下：").bold = True
        docs.add_paragraph().add_run("一审判决认定事实清楚，适用法律正确，请求法院驳回上诉人上诉请求，维持原判。理由如下：").bold = True

        for i in (conten[4]+"\n综上，原告诉讼请求缺乏事实与法律依据，答辩人请求法院依法裁决。\n此致\n"+conten[13][0][0]).split("\n"):
            docs.add_paragraph().add_run(i)

        p = docs.add_paragraph()
        p.add_run("答辩人：")
        p.paragraph_format.left_indent = Pt(252)
        p = docs.add_paragraph()
        p.add_run("年    月    日")
        p.alignment = WD_TABLE_ALIGNMENT.RIGHT

        docs.add_paragraph()

        docs.add_paragraph().add_run("附：答辩状副本   份")

        ps = docs.paragraphs
        for i in ps[1:]:
            format = i.paragraph_format
            if i != ps[-5]:
                format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    try:
        if not os.path.exists(dir):
            run('mkdir "{}"'.format(dir),shell=True)
        file = dir + "\\" + str(time.strftime("%Y%m%d%H%M%S")) + "-"+what+".docx"
        docs.save(file)
        os.startfile(file)
    except:
        pass

output_paper()

