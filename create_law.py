import os.path
import re
import time

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


def create_law(files,sig):
    law_txts = ""
    for i in files:
        i = i.replace("/","\\")
        f_e = os.path.splitext(i)
        if i.endswith(".txt"):
            with open(i, "r", encoding="utf-8-sig") as t:
                codes = t.readlines()
        else:
            try:
                docs = docx.Document(i)
                codes = [i.text for i in docs.paragraphs]
            except:
                continue
        ok = kk = 0
        zjp = []
        txt = "laws\\"+f_e[0].split("\\")[-1].replace("中华人民共和国","").replace("人民代表大会","人大").replace("常务委员会","常委")+".txt"
        if os.path.exists(txt):
            os.remove(txt)
        law_txt = []
        with open(txt,"a",encoding="utf-8") as tf:
            for j in codes:
                t = re.sub("\s+"," ",re.sub("^\s+","",j)) + "\n"
                c = re.findall("^第.+?([编章节])[、.\s]+", t)
                if kk == 0 and len(c) > 0:
                    if c[0] not in zjp:
                        zjp.append(c[0])
                        tf.write(t)
                        law_txt.append(t)
                        continue
                    else:
                        kk = 1
                if re.search("(^第.+?条[、.\s]+)|(^[0-9一二三四五六七八九十][、.\s]+)",t) is not None:
                    ok = 1
                if ok == 1:
                    tf.write(t)
                    law_txt.append(t)

        if len(law_txt) == 0:
            os.remove(txt)
            law_txts = law_txts + f_e[0]+"\n"

    sig.emit(law_txts)

def output_paper(dir,conten,what):
    docs = docx.Document()
    p = docs.add_paragraph()
    p.alignment = WD_TABLE_ALIGNMENT.CENTER

    format = p.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(0)

    run = p.add_run(what)
    run.font.size = Pt(18)
    run.font.name = u'黑体'
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    docs.add_paragraph()

    if what == "民事起诉状":
        for i in ["原告","被告","第三人"]:
            for j in conten[i]:
                if j[0] == "":
                    continue
                p = docs.add_paragraph()
                p.add_run(i+"：").bold = True
                p.add_run(j[0]+"，"+j[1].split("\n")[0])
                for k in j[1].split("\n")[1:]:
                    docs.add_paragraph().add_run(k)

        docs.add_paragraph()
        p = docs.add_paragraph()
        p.add_run("案由：").bold = True
        p.add_run(conten["案由"][0][0])
        docs.add_paragraph()

        docs.add_paragraph().add_run("诉讼请求：").bold = True
        n = 0
        for i in conten["诉讼请求"]:
            n = n + 1
            docs.add_paragraph(str(n)+". "+i[1])

        docs.add_paragraph().add_run("事实理由：").bold = True
        for i in [*(conten["事实理由"].split("\n")),"现为维护自身的权利，原告根据《中华人民共和国民事诉讼法》及相关规定，向贵院提起诉讼，请法院判如所请。","此致",conten["管辖法院"][0]]:
            docs.add_paragraph().add_run(i)

        p = docs.add_paragraph()
        p.add_run("具状人：")
        p.paragraph_format.left_indent = Pt(252)
        p = docs.add_paragraph()
        p.add_run("年    月    日")
        p.alignment = WD_TABLE_ALIGNMENT.RIGHT

        docs.add_paragraph()

        docs.add_paragraph().add_run("附：1. 起诉状副本   份")
        for i in ["2. 证据：","1)原告身份证复印件/营业执照复印件、法定代表人身份证明、法定代表人身份证复印件；","2)被告身份证复印件/身份信息/营业执照复印件/企业信用信息打印件；"]:
            p = docs.add_paragraph()
            p.add_run(i)
            p.paragraph_format.left_indent = Pt(28)

        ps = docs.paragraphs
        for i in ps[1:]:
            format = i.paragraph_format
            if i not in [ps[-8],ps[-2:]]:
                format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    elif what == "工作日志":
        for i in conten:
            docs.add_paragraph().add_run(i[1].replace("-",".") +"-"+i[7].replace("-","."))
            docs.add_paragraph().add_run(i[3])
            for j in range(4,7):
                if i[j] != "":
                    docs.add_paragraph().add_run(i[j])
            docs.add_paragraph()

        for i in docs.paragraphs[1:]:
            format = i.paragraph_format
            format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    else:
        other = []
        client = ["","",""]
        for i in ["原告","被告","第三人"]:
            for j in conten[i]:
                if "委托人" in j:
                    client = j
                else:
                    other.append(j)
        p = docs.add_paragraph()
        p.add_run("答辩人：").bold = True
        p.add_run(client[0] + "，" + client[1].split("\n")[0])
        for i in client[1].split("\n")[1:]:
            docs.add_paragraph().add_run(i)

        docs.add_paragraph()
        docs.add_paragraph().add_run("答辩人与"+"、".join([i[0] for i in other])[0:-1]+conten["案由"][0][0]+"一案，针对原告/上诉人/申请人的起诉/上诉/仲裁理由答辩如下：")

        docs.add_paragraph().add_run("原告的诉讼请求缺乏事实与法律依据，请求法院驳回原告诉讼请求。理由如下：").bold = True
        docs.add_paragraph().add_run("一审判决认定事实清楚，适用法律正确，请求法院驳回上诉人上诉请求，维持原判。理由如下：").bold = True

        for i in (conten["事实理由"]+"\n综上，原告诉讼请求缺乏事实与法律依据，答辩人请求法院依法裁决。\n此致\n"+conten["管辖法院"][0][0]).split("\n"):
            docs.add_paragraph().add_run(i)

        p = docs.add_paragraph()
        p.add_run("答辩人：")
        p.paragraph_format.left_indent = Pt(252)
        p = docs.add_paragraph()
        p.add_run("年    月    日")
        p.alignment = WD_TABLE_ALIGNMENT.RIGHT

        docs.add_paragraph()

        docs.add_paragraph().add_run("附：答辩状副本   份")

        ps = docs.paragraphs
        for i in ps[1:]:
            format = i.paragraph_format
            if i != ps[-5]:
                format.first_line_indent = Pt(28)
            format.space_before = Pt(0)
            format.space_after = Pt(0)
            format.line_spacing = Pt(24)
            for j in i.runs:
                j.font.size = Pt(14)
                j.font.name = u'仿宋'
                j.font.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    try:
        docs.save(dir + "\\" + str(time.strftime("%Y%m%d%H%M%S")) + "-"+what+".docx")
        os.startfile(dir + "\\" + str(time.strftime("%Y%m%d%H%M%S")) + "-"+what+".docx")
    except:
        pass

